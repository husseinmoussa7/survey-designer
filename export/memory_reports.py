"""
export/memory_reports.py
For each training memory file: exports a .md and .docx showing
hypothesis, survey_v0, meeting discussion, critic output, survey_v1.
Full question options (scale labels, choices, matrix statements) are included.
Output: results/memory_reports/{project_name}.md  /  .docx

Usage:
  python export/memory_reports.py
"""
import json
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MEMORY    = Path("memory")
PAPERS_DIR = Path(os.getenv("PAPERS_DIR", "papers"))
OUT_DIR   = Path("results/memory_reports")
OUT_DIR.mkdir(parents=True, exist_ok=True)

SEV_COLOR = {3: RGBColor(0xC0,0x00,0x00), 2: RGBColor(0xCC,0x66,0x00), 1: RGBColor(0x20,0x6B,0xC0)}
SEV_LABEL = {3: "MUST FIX", 2: "MODERATE", 1: "MINOR"}

DETAIL_PAPERS = {"01_olson_2023_commoncents_study_1",
                 "05_ding_2025_animationspeedconvex_study_1b"}


# ── Papers directory lookup ───────────────────────────────────────────────────
def find_hypotheses(project_name: str) -> str:
    """Return hypotheses.txt content for the given project, or empty string."""
    num = project_name.split("_")[0].zfill(2)
    for subdir in ("training", "test", ""):
        base = PAPERS_DIR / subdir if subdir else PAPERS_DIR
        if not base.exists():
            continue
        paper_dir = next((d for d in base.iterdir()
                          if d.is_dir() and d.name.split("_")[0].zfill(2) == num), None)
        if not paper_dir:
            continue
        all_hyp = sorted(paper_dir.rglob("hypotheses.txt"))
        if not all_hyp:
            continue
        best = all_hyp[0]
        for hyp in all_hyp:
            study = hyp.parent.name.lower().replace(" ", "_")
            if project_name.endswith(study):
                best = hyp
                break
        try:
            return best.read_text(encoding="utf-8", errors="ignore").strip()
        except Exception:
            return ""
    return ""


# ── Helpers ───────────────────────────────────────────────────────────────────
def set_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)

def add_border(table, color="AAAAAA"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    bords = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color); bords.append(el)
    tblPr.append(bords)


# ── Full question rendering — Markdown ────────────────────────────────────────
def render_question_md(q, index, prefix=""):
    qid   = q.get("question_id", "?")
    qtext = q.get("question_text", "")
    qtype = q.get("input_type", "?").lower()
    cfg   = q.get("input_config", {})
    lines = [f"{prefix}**Q{index}. [{qid}]** {qtext}"]

    if qtype in ("scale", "likert"):
        mn = cfg.get("scale_min", cfg.get("min_value", 1))
        mx = cfg.get("scale_max", cfg.get("scale_points", cfg.get("max_value", 7)))
        labels = cfg.get("scale_labels", [])
        min_lbl = cfg.get("scale_min_label", "")
        max_lbl = cfg.get("scale_max_label", "")
        if labels:
            label_str = "  ".join(f"{mn+i}={l}" for i, l in enumerate(labels))
            lines.append(f"{prefix}*Scale {mn}–{mx}:* {label_str}")
        elif min_lbl or max_lbl:
            lines.append(f"{prefix}*Scale {mn}–{mx}:* {min_lbl} → {max_lbl}")
        else:
            lines.append(f"{prefix}*Scale {mn}–{mx}*")
    elif qtype in ("multiple_choice", "single_choice", "radio"):
        opts = cfg.get("options", [])
        lines.append(f"{prefix}*{'Multiple' if 'multiple' in qtype else 'Single'} choice:*")
        for opt in opts:
            lines.append(f"{prefix}  - {opt}")
        if cfg.get("allow_other"):
            lines.append(f"{prefix}  - Other (please specify)")
    elif qtype in ("matrix", "matrix_scale"):
        stmts  = cfg.get("statements", cfg.get("rows", []))
        labels = cfg.get("scale_labels", [])
        mn     = cfg.get("scale_min", 1)
        mx     = cfg.get("scale_max", 7)
        scale_str = f"Scale {mn}–{mx}" + (f": {' | '.join(labels)}" if labels else "")
        lines.append(f"{prefix}*Matrix ({scale_str}):*")
        for stmt in stmts:
            lines.append(f"{prefix}  - {stmt}")
    elif qtype in ("open_text", "text_input", "textarea"):
        ml = cfg.get("max_length", "")
        lines.append(f"{prefix}*Open text{(' — max ' + str(ml) + ' chars') if ml else ''}*")
    elif qtype == "number":
        mn2 = cfg.get("min_value", "")
        mx2 = cfg.get("max_value", "")
        lines.append(f"{prefix}*Numeric input{(' ('+str(mn2)+'–'+str(mx2)+')') if mn2 or mx2 else ''}*")
    else:
        lines.append(f"{prefix}*{qtype}*")

    lines.append("")
    return "\n".join(lines)


def render_questions_md(questions, prefix=""):
    return "\n".join(render_question_md(q, i, prefix) for i, q in enumerate(questions, 1))


# ── Full question rendering — Word ────────────────────────────────────────────
def render_question_docx(doc, q, index, is_new=False):
    qid   = q.get("question_id", "?")
    qtext = q.get("question_text", "")
    qtype = q.get("input_type", "?").lower()
    cfg   = q.get("input_config", {})
    green = RGBColor(0x1F, 0x75, 0x23)

    p = doc.add_paragraph()
    r = p.add_run(f"Q{index}. [{qid}]  ")
    r.bold = True
    if is_new: r.font.color.rgb = green
    body = p.add_run(qtext)
    body.font.size = Pt(11)
    if is_new: body.bold = True

    if is_new:
        tag = doc.add_paragraph()
        run = tag.add_run("  ★ NEW — added by critic-driven revision")
        run.italic = True; run.font.size = Pt(9); run.font.color.rgb = green

    if qtype in ("scale", "likert"):
        mn = cfg.get("scale_min", cfg.get("min_value", 1))
        mx = cfg.get("scale_max", cfg.get("scale_points", cfg.get("max_value", 7)))
        labels = cfg.get("scale_labels", [])
        min_lbl = cfg.get("scale_min_label", "")
        max_lbl = cfg.get("scale_max_label", "")
        if labels:
            n = len(labels)
            tbl = doc.add_table(rows=2, cols=n)
            tbl.style = "Table Grid"
            add_border(tbl, "CCCCCC")
            for j, lbl in enumerate(labels):
                nc = tbl.rows[0].cells[j]
                nc.text = str(mn + j)
                nc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                nc.paragraphs[0].runs[0].font.size = Pt(9)
                nc.paragraphs[0].runs[0].bold = True
                set_bg(nc, "EAF4EA" if is_new else "EBF3FB")
                lc = tbl.rows[1].cells[j]
                lc.text = lbl
                lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                lc.paragraphs[0].runs[0].font.size = Pt(8)
        elif min_lbl or max_lbl:
            sp = doc.add_paragraph()
            sp.add_run(f"Scale {mn}–{mx}:  ").bold = True
            sp.add_run(f"{min_lbl}  →  {max_lbl}").italic = True
            sp.runs[-1].font.size = Pt(10)
        else:
            sp = doc.add_paragraph()
            sp.add_run(f"Scale {mn}–{mx}").italic = True
            sp.runs[0].font.size = Pt(10)
        doc.add_paragraph()
    elif qtype in ("multiple_choice", "single_choice", "radio"):
        opts = cfg.get("options", [])
        for opt in opts:
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run(opt).font.size = Pt(10)
        if cfg.get("allow_other"):
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run("Other (please specify)").font.size = Pt(10)
            bp.runs[0].italic = True
    elif qtype in ("matrix", "matrix_scale"):
        stmts  = cfg.get("statements", cfg.get("rows", []))
        labels = cfg.get("scale_labels", [])
        mn     = cfg.get("scale_min", 1)
        if stmts and labels:
            ncols = len(labels) + 1
            tbl   = doc.add_table(rows=len(stmts) + 1, cols=ncols)
            tbl.style = "Table Grid"
            add_border(tbl, "CCCCCC")
            set_bg(tbl.rows[0].cells[0], "F2F2F2")
            for j, lbl in enumerate(labels):
                c = tbl.rows[0].cells[j + 1]
                c.text = lbl
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                c.paragraphs[0].runs[0].font.size = Pt(8)
                c.paragraphs[0].runs[0].bold = True
                set_bg(c, "EAF4EA" if is_new else "EBF3FB")
            for ri, stmt in enumerate(stmts):
                tbl.rows[ri + 1].cells[0].text = stmt
                tbl.rows[ri + 1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                for j in range(len(labels)):
                    c = tbl.rows[ri + 1].cells[j + 1]
                    c.text = "○"
                    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stmts:
            for stmt in stmts:
                bp = doc.add_paragraph(style="List Bullet")
                bp.add_run(stmt).font.size = Pt(10)
        doc.add_paragraph()
    elif qtype in ("open_text", "text_input", "textarea"):
        ml = cfg.get("max_length", "")
        op = doc.add_paragraph()
        r  = op.add_run("[ Open response ]")
        r.italic = True; r.font.color.rgb = RGBColor(0x88,0x88,0x88); r.font.size = Pt(10)
        if ml:
            op.add_run(f"  (max {ml} chars)").font.size = Pt(9)
    elif qtype == "number":
        op = doc.add_paragraph()
        op.add_run("[ Numeric input ]").italic = True
        op.runs[0].font.color.rgb = RGBColor(0x88,0x88,0x88)
        op.runs[0].font.size = Pt(10)
    else:
        op = doc.add_paragraph()
        op.add_run(f"[ {qtype} ]").italic = True
        op.runs[0].font.size = Pt(10)

    doc.add_paragraph()


# ── Critique normalization ────────────────────────────────────────────────────
ISSUE_KEYS = {"leading_language","double_barreled","social_desirability_bias",
              "ambiguous_wording","scale_appropriateness"}

def normalize_critique(crit: dict):
    q_crits = crit.get("question_critiques", [])
    s_crits_raw = crit.get("survey_level_issues", [])
    issue_map = {}
    for entry in q_crits:
        qid = entry.get("question_id", "?")
        if "issues" in entry:
            issues = entry["issues"]
        elif any(k in entry for k in ISSUE_KEYS):
            issues = []
            for k in ISSUE_KEYS:
                if k in entry:
                    sev = entry[k]
                    if isinstance(sev, int) and sev > 0:
                        issues.append({"type": k.replace("_"," "), "severity": sev,
                                       "explanation": entry.get("comments","")})
        else:
            issues = []
        issue_map[qid] = [i for i in issues if isinstance(i, dict)]
    s_crits = []
    for item in s_crits_raw:
        if isinstance(item, dict):
            if "type" in item:
                s_crits.append(item)
            else:
                for k, v in item.items():
                    if isinstance(v, str) and v:
                        s_crits.append({"type": k.replace("_"," "), "severity": 2, "explanation": v})
        elif isinstance(item, str):
            s_crits.append({"type": "issue", "severity": 2, "explanation": item})
    return issue_map, [i for i in s_crits if isinstance(i, dict)]


# ── Per-file generator ────────────────────────────────────────────────────────
def process_file(path: Path, write_detail: bool):
    data    = json.loads(path.read_text())
    project = data.get("project", path.stem)
    ts      = data.get("timestamp","")[:16].replace("T"," ")
    disc    = data.get("discussion_log", [])
    crit    = data.get("critique", {}) if isinstance(data.get("critique"), dict) else {}
    v0      = data.get("survey_v0", {})
    v1      = data.get("survey_v1", {})

    v0_qs   = v0.get("revised_survey", v0).get("questions", []) if v0 else []
    v1_qs   = v1.get("revised_survey", v1).get("questions", []) if v1 else []
    v0_ids  = {q["question_id"] for q in v0_qs}

    issue_map, s_crits = normalize_critique(crit)
    all_issues_flat = [i for iss in issue_map.values() for i in iss]
    q_crits = crit.get("question_critiques", [])
    new_ids = [q["question_id"] for q in v1_qs if q["question_id"] not in v0_ids]

    hypotheses = find_hypotheses(project)

    # ── Markdown ──────────────────────────────────────────────────────────────
    md = f"# Pipeline Report: {project}\n"
    md += f"*Run: {ts}*\n\n---\n\n"

    md += "## Research Hypotheses\n\n"
    if hypotheses:
        for line in hypotheses.splitlines():
            md += f"{line}\n"
    else:
        md += "*Hypotheses not found.*\n"
    md += "\n---\n\n"

    md += f"## Survey V0 — Original Converted Survey ({len(v0_qs)} questions)\n\n"
    md += render_questions_md(v0_qs) + "\n---\n\n"

    md += f"## Meeting Discussion ({len(disc)//2} rounds)\n\n"
    if not disc:
        md += "*No discussion log recorded.*\n\n"
    else:
        for entry in disc:
            role    = entry.get("role","?").upper()
            round_  = entry.get("round","?")
            content = entry.get("content","")
            md += f"### Round {round_} — {role}\n\n{content}\n\n"
    md += "---\n\n"

    md += "## Critic Output\n\n"
    md += f"**{len(all_issues_flat)} question-level issues** across {len(q_crits)} questions  |  "
    md += f"**{len(s_crits)} survey-level issues**\n\n"

    md += "### Question-Level Issues\n\n"
    for entry in q_crits:
        qid    = entry.get("question_id","?")
        issues = issue_map.get(qid,[])
        if not issues: continue
        md += f"**{qid}:**\n"
        for iss in issues:
            sev = iss.get("severity",0)
            md += f"- [{SEV_LABEL.get(sev,f'SEV {sev}')}] **{iss.get('type','?')}**: {iss.get('explanation','')}\n"
        md += "\n"

    md += "### Survey-Level Issues\n\n"
    for iss in s_crits:
        sev = iss.get("severity",0)
        md += f"- [{SEV_LABEL.get(sev,f'SEV {sev}')}] **{iss.get('type','?')}**: {iss.get('explanation','')}\n"
    md += "\n---\n\n"

    md += f"## Survey V1 — Final Revised Survey ({len(v1_qs)} questions)\n\n"
    if new_ids:
        md += f"*New questions added: {', '.join(new_ids)}*\n\n"
    md += render_questions_md(v1_qs) + "\n---\n\n"

    if write_detail:
        md += "## Before / After: Issues and Fixes\n\n"
        for q in v1_qs:
            qid    = q["question_id"]
            issues = issue_map.get(qid,[])
            before = next((x for x in v0_qs if x["question_id"]==qid), None)
            is_new = qid not in v0_ids
            if not issues and not is_new: continue
            md += f"### {qid}\n\n"
            if is_new:
                md += f"**★ NEW** — Added to address survey-level gaps\n\n"
                md += f"*Question:* {q.get('question_text','')}\n\n"
            else:
                md += f"**Before:** {before.get('question_text','') if before else '—'}\n\n"
                md += f"**After:** {q.get('question_text','')}\n\n"
                for iss in issues:
                    sev = iss.get("severity",0)
                    md += f"- [{SEV_LABEL.get(sev,f'SEV {sev}')}] {iss.get('type','?')}: {iss.get('explanation','')}\n"
                md += "\n"
        md += "---\n\n"

    md += "*Generated by Survey Designer Agent*\n"
    (OUT_DIR / f"{project}.md").write_text(md)

    # ── Word doc ──────────────────────────────────────────────────────────────
    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1.0)

    doc.add_heading(f"Pipeline Report: {project}", level=1)
    doc.add_paragraph(f"Run: {ts}").runs[0].italic = True
    doc.add_paragraph()

    doc.add_heading("Research Hypotheses", level=2)
    if hypotheses:
        for line in hypotheses.splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())
    else:
        doc.add_paragraph("Hypotheses not found.").runs[0].italic = True
    doc.add_paragraph()

    doc.add_heading(f"Survey V0 — Original Converted Survey  ({len(v0_qs)} questions)", level=2)
    for i, q in enumerate(v0_qs, 1):
        render_question_docx(doc, q, i, is_new=False)
    doc.add_paragraph()

    doc.add_heading(f"Meeting Discussion  ({len(disc)//2} rounds)", level=2)
    if not disc:
        doc.add_paragraph("No discussion log recorded.").runs[0].italic = True
    else:
        for entry in disc:
            role    = entry.get("role","?").title()
            round_  = entry.get("round","?")
            content = entry.get("content","")
            role_color = (RGBColor(0xC0,0x00,0x00) if role.lower()=="critic"
                          else RGBColor(0x7B,0x2F,0xBE) if role.lower()=="econometrician"
                          else RGBColor(0x1F,0x75,0x23))
            h = doc.add_heading(f"Round {round_} — {role}", level=3)
            h.runs[0].font.color.rgb = role_color
            doc.add_paragraph(content)
    doc.add_paragraph()

    doc.add_heading("Critic Output", level=2)
    p = doc.add_paragraph()
    p.add_run(f"{len(all_issues_flat)} question-level issues").bold = True
    p.add_run(f"  across {len(q_crits)} questions   |   ")
    p.add_run(f"{len(s_crits)} survey-level issues").bold = True

    doc.add_heading("Question-Level Issues", level=3)
    for entry in q_crits:
        qid    = entry.get("question_id","?")
        issues = issue_map.get(qid,[])
        if not issues: continue
        doc.add_paragraph(qid, style="List Bullet").runs[0].bold = True
        for iss in issues:
            sev   = iss.get("severity",0)
            color = SEV_COLOR.get(sev, RGBColor(0,0,0))
            label = SEV_LABEL.get(sev, f"SEV {sev}")
            bp    = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.left_indent = Inches(0.3)
            tag   = bp.add_run(f"[{label}]  ")
            tag.bold = True; tag.font.color.rgb = color
            typ   = bp.add_run(f"{iss.get('type','?')}:  ")
            typ.bold = True
            bp.add_run(iss.get("explanation","")).font.size = Pt(10)

    doc.add_heading("Survey-Level Issues", level=3)
    for iss in s_crits:
        sev   = iss.get("severity",0)
        color = SEV_COLOR.get(sev, RGBColor(0,0,0))
        label = SEV_LABEL.get(sev, f"SEV {sev}")
        bp    = doc.add_paragraph(style="List Bullet")
        tag   = bp.add_run(f"[{label}]  ")
        tag.bold = True; tag.font.color.rgb = color
        typ   = bp.add_run(f"{iss.get('type','?')}:  ")
        typ.bold = True
        bp.add_run(iss.get("explanation","")).font.size = Pt(10)
    doc.add_paragraph()

    doc.add_heading(f"Survey V1 — Final Revised Survey  ({len(v1_qs)} questions)", level=2)
    if new_ids:
        p = doc.add_paragraph()
        p.add_run("New questions added: ").bold = True
        p.add_run(", ".join(new_ids))
    for i, q in enumerate(v1_qs, 1):
        is_new = q["question_id"] not in v0_ids
        render_question_docx(doc, q, i, is_new=is_new)
    doc.add_paragraph()

    if write_detail:
        doc.add_heading("Before / After: Issues and Fixes", level=2)
        for q in v1_qs:
            qid    = q["question_id"]
            issues = issue_map.get(qid,[])
            before = next((x for x in v0_qs if x["question_id"]==qid), None)
            is_new = qid not in v0_ids
            if not issues and not is_new: continue
            doc.add_heading(qid, level=3)
            if is_new:
                p = doc.add_paragraph()
                p.add_run("★ NEW").bold = True
                p.add_run(" — Added to address survey-level gaps")
                doc.add_paragraph(q.get("question_text","")).runs[0].italic = True
            else:
                tbl = doc.add_table(rows=2, cols=2)
                tbl.style = "Table Grid"
                add_border(tbl)
                set_bg(tbl.rows[0].cells[0], "FFE7E7")
                set_bg(tbl.rows[0].cells[1], "E2EFDA")
                tbl.rows[0].cells[0].text = "BEFORE"
                tbl.rows[0].cells[1].text = "AFTER"
                for c in tbl.rows[0].cells:
                    c.paragraphs[0].runs[0].bold = True
                    c.paragraphs[0].runs[0].font.size = Pt(9)
                tbl.rows[1].cells[0].text = before.get("question_text","—") if before else "—"
                tbl.rows[1].cells[1].text = q.get("question_text","")
                for c in tbl.rows[1].cells:
                    c.paragraphs[0].runs[0].font.size = Pt(10)
                doc.add_paragraph()
                for iss in issues:
                    sev   = iss.get("severity",0)
                    color = SEV_COLOR.get(sev, RGBColor(0,0,0))
                    label = SEV_LABEL.get(sev, f"SEV {sev}")
                    bp    = doc.add_paragraph(style="List Bullet")
                    tag   = bp.add_run(f"[{label}]  ")
                    tag.bold = True; tag.font.color.rgb = color
                    bp.add_run(f"{iss.get('type','?')}: {iss.get('explanation','')}").font.size = Pt(9)
            doc.add_paragraph()

    doc.add_paragraph("Generated by Survey Designer Agent").runs[0].italic = True

    docx_path = OUT_DIR / f"{project}.docx"
    doc.save(docx_path)
    return OUT_DIR / f"{project}.md", docx_path


# ── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    skip = {"hotel_booking_study", "shani_feinstein_2022"}
    files = sorted(MEMORY.glob("*.json"))

    latest = {}
    for f in files:
        proj = "_".join(f.stem.split("_")[:-2])
        if proj not in skip and not any(s in proj for s in skip):
            if proj not in latest or f.stat().st_mtime > latest[proj].stat().st_mtime:
                latest[proj] = f

    print(f"Exporting {len(latest)} memory files...")
    for proj, path in sorted(latest.items()):
        detail = proj in DETAIL_PAPERS
        md_p, docx_p = process_file(path, write_detail=detail)
        tag = "  [+detail]" if detail else ""
        print(f"  {proj}{tag}")

    print(f"\nAll reports saved to: {OUT_DIR}/")
    print(f"  .md  files: {len(list(OUT_DIR.glob('*.md')))}")
    print(f"  .docx files: {len(list(OUT_DIR.glob('*.docx')))}")
