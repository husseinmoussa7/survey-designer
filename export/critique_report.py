"""
export/critique_report.py
Produces critique_with_revisions.docx — a side-by-side view of:
  - Original survey (v0) vs final revised survey (v1)
  - Critic's issues per question
  - Added questions with explanation

Usage:
  python export/critique_report.py
  python export/critique_report.py --project my_project_name
"""
import argparse
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DEFAULT_PROJECT = "shani_feinstein_2022"

SEV_COLOR  = {3: RGBColor(0xC0, 0x00, 0x00), 2: RGBColor(0xFF, 0x80, 0x00), 1: RGBColor(0x00, 0x70, 0xC0)}
SEV_LABEL  = {3: "MUST FIX", 2: "MODERATE", 1: "MINOR"}
SEV_BG     = {3: "FFE7E7", 2: "FFF3E0", 1: "E8F0FE"}


def add_border(table):
    tbl   = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    bords = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "AAAAAA")
        bords.append(el)
    tblPr.append(bords)


def set_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, color=None, size=10, italic=False):
    cell.text = ""
    p    = cell.paragraphs[0]
    run  = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def q_summary(q):
    t   = q.get("input_type", "")
    cfg = q.get("input_config", {})
    if t == "scale":
        mn = cfg.get("scale_min", 1)
        mx = cfg.get("scale_max", 7)
        labels = [l for l in cfg.get("scale_labels", []) if l]
        ends   = f"{labels[0]} → {labels[-1]}" if len(labels) >= 2 else ""
        return f"Scale {mn}–{mx}  ({ends})"
    if t == "multiple_choice":
        opts = cfg.get("options", [])
        return "Choices: " + " | ".join(opts[:3]) + ("…" if len(opts) > 3 else "")
    if t == "matrix":
        stmts = cfg.get("statements", [])
        return f"Matrix — {len(stmts)} statement(s), scale {cfg.get('scale_min',1)}–{cfg.get('scale_max',7)}"
    if t == "open_text":
        return f"Open text (max {cfg.get('max_length','?')} chars)"
    return t


def build_doc(project_name: str):
    dir_path  = Path("results") / project_name
    mem_files = sorted(Path("memory").glob(f"{project_name}_*.json"))
    if not mem_files:
        raise FileNotFoundError(f"No memory file found for project '{project_name}'")

    mem        = json.loads(mem_files[-1].read_text())
    v0_qs      = mem["survey_v0"].get("revised_survey", mem["survey_v0"]).get("questions", [])
    final_data = json.loads((dir_path / "ai_survey.json").read_text())
    final_qs   = final_data.get("revised_survey", final_data).get("questions", [])
    explanations = final_data.get("explanations", {})
    critique   = json.loads((dir_path / "critique.json").read_text())

    v0_map    = {q["question_id"]: q for q in v0_qs}
    final_map = {q["question_id"]: q for q in final_qs}

    issue_map = {}
    for entry in critique.get("question_critiques", []):
        qid    = entry.get("question_id", "?")
        issues = entry.get("issues", [])
        if not issues and "severity" in entry:
            issues = [entry]
        issue_map[qid] = issues
    survey_issues = critique.get("survey_level_issues", [])

    doc = Document()
    doc.core_properties.title = "Survey Critique + Revisions"
    for section in doc.sections:
        section.left_margin  = Inches(0.9)
        section.right_margin = Inches(0.9)

    h = doc.add_heading("Survey: Critique & Revisions", level=1)
    doc.add_paragraph(f"{project_name}").runs[0].italic = True
    doc.add_paragraph(
        f"Original survey: {len(v0_qs)} questions  →  "
        f"Final survey: {len(final_qs)} questions  "
        f"(+{len(final_qs)-len(v0_qs)} added by critic-driven revision)"
    )
    doc.add_paragraph()

    # Part 1: Question-by-question
    doc.add_heading("Part 1 — Question-Level: Before / Critique / After", level=1)

    all_qids = list(dict.fromkeys(list(v0_map.keys()) + list(final_map.keys())))

    for qid in all_qids:
        before  = v0_map.get(qid)
        after   = final_map.get(qid)
        issues  = issue_map.get(qid, [])
        is_new  = qid not in v0_map
        explain = explanations.get(qid, "")

        status = " ✦ NEW" if is_new else ""
        doc.add_heading(f"Question {qid.upper()}{status}", level=2)

        tbl = doc.add_table(rows=3, cols=2)
        tbl.style = "Table Grid"
        add_border(tbl)
        set_bg(tbl.rows[0].cells[0], "D9E1F2")
        set_bg(tbl.rows[0].cells[1], "E2EFDA")
        cell_text(tbl.rows[0].cells[0], "BEFORE  (original / pre-enhancement)", bold=True, size=10)
        cell_text(tbl.rows[0].cells[1], "AFTER   (final, post-critique revision)", bold=True, size=10)

        before_text = before["question_text"] if before else "—  (question did not exist)"
        after_text  = after["question_text"]  if after  else "—  (question removed)"
        cell_text(tbl.rows[1].cells[0], before_text, size=10)
        cell_text(tbl.rows[1].cells[1], after_text,  size=10)

        before_cfg = q_summary(before) if before else "n/a"
        after_cfg  = q_summary(after)  if after  else "n/a"
        cell_text(tbl.rows[2].cells[0], before_cfg, italic=True, color=RGBColor(0x60,0x60,0x60), size=9)
        cell_text(tbl.rows[2].cells[1], after_cfg,  italic=True, color=RGBColor(0x60,0x60,0x60), size=9)
        doc.add_paragraph()

        if is_new:
            p = doc.add_paragraph()
            p.add_run("Why this question was added:  ").bold = True
            p.add_run(explain or "Added by critic-driven revision to address survey-level gaps.")
        elif not issues:
            p = doc.add_paragraph()
            p.add_run("No issues flagged by critic for this question.").italic = True
        else:
            doc.add_paragraph().add_run("Critic issues:").bold = True
            for iss in issues:
                sev   = iss.get("severity", 0)
                itype = iss.get("type", "?")
                expl  = iss.get("explanation", "")
                color = SEV_COLOR.get(sev, RGBColor(0,0,0))
                label = SEV_LABEL.get(sev, f"SEV {sev}")
                p   = doc.add_paragraph(style="List Bullet")
                tag = p.add_run(f"[{label}]  ")
                tag.bold = True
                tag.font.color.rgb = color
                typ = p.add_run(f"{itype}:  ")
                typ.bold = True
                p.add_run(expl)

        doc.add_paragraph()
        doc.add_paragraph("─" * 80).runs[0].font.color.rgb = RGBColor(0xCC,0xCC,0xCC)
        doc.add_paragraph()

    # Part 2: Survey-level issues
    doc.add_heading("Part 2 — Survey-Level Issues & How They Were Addressed", level=1)
    doc.add_paragraph(
        "These issues apply to the survey as a whole. The critic flagged them in Step 3; "
        "the Step 4 revision pass added new questions to address them."
    )
    doc.add_paragraph()
    for iss in survey_issues:
        sev   = iss.get("severity", 0)
        itype = iss.get("type", "?")
        expl  = iss.get("explanation", "")
        color = SEV_COLOR.get(sev, RGBColor(0,0,0))
        label = SEV_LABEL.get(sev, f"SEV {sev}")
        tbl = doc.add_table(rows=2, cols=1)
        tbl.style = "Table Grid"
        add_border(tbl)
        set_bg(tbl.rows[0].cells[0], SEV_BG.get(sev, "FFFFFF"))
        cell_text(tbl.rows[0].cells[0], f"[{label}]  {itype.upper()}", bold=True, color=color, size=11)
        cell_text(tbl.rows[1].cells[0], expl, size=10)
        doc.add_paragraph()
        addressed = [qid for qid in final_map if qid not in v0_map]
        if addressed and sev >= 2:
            p = doc.add_paragraph()
            p.add_run("Addressed by new questions: ").bold = True
            p.add_run(", ".join(addressed))
        doc.add_paragraph()

    # Part 3: Summary table
    doc.add_heading("Part 3 — Summary: What Changed", level=1)
    new_qids  = [qid for qid in final_map if qid not in v0_map]
    kept_qids = [qid for qid in v0_map if qid in final_map]
    total_iss  = sum(len(v) for v in issue_map.values())
    sev3_q     = sum(1 for v in issue_map.values() for i in v if i.get("severity") == 3)
    sev3_s     = sum(1 for i in survey_issues if i.get("severity") == 3)

    rows = [
        ("Original questions (v0)",        str(len(v0_qs))),
        ("Final questions (post-critique)", str(len(final_qs))),
        ("Questions added by revision",     str(len(new_qids)) + f"  ({', '.join(new_qids)})"),
        ("Questions retained unchanged",    str(len(kept_qids))),
        ("Total critic issues (question-level)", str(total_iss)),
        ("Severity-3 question issues",      str(sev3_q)),
        ("Severity-3 survey-level issues",  str(sev3_s)),
        ("Step 4 revision triggered?",      "Yes" if (sev3_q + sev3_s) > 0 else "No"),
    ]
    tbl = doc.add_table(rows=len(rows)+1, cols=2)
    tbl.style = "Table Grid"
    add_border(tbl)
    set_bg(tbl.rows[0].cells[0], "D9E1F2"); set_bg(tbl.rows[0].cells[1], "D9E1F2")
    cell_text(tbl.rows[0].cells[0], "Metric", bold=True)
    cell_text(tbl.rows[0].cells[1], "Value",  bold=True)
    for i, (k, v) in enumerate(rows):
        cell_text(tbl.rows[i+1].cells[0], k, size=10)
        cell_text(tbl.rows[i+1].cells[1], v, size=10)

    doc.add_paragraph()
    doc.add_paragraph("Generated by Survey Designer Agent").runs[0].italic = True

    out = dir_path / "critique_with_revisions.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Export critique + revisions report")
    parser.add_argument("--project", default=DEFAULT_PROJECT,
                        help=f"Project name (default: {DEFAULT_PROJECT})")
    args = parser.parse_args()
    build_doc(args.project)
