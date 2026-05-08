"""
scripts/evaluate.py
Evaluates AI-generated surveys against ground truth using GPT-4o as judge.

Metrics:
  1. Question count & type distribution
  2. Construct coverage vs hypotheses (+ vs ground truth if available)
  3. Wording quality (GPT-4o scores each question 1–5)
  4. Scale consistency
  5. Critic issue summary

Usage:
  python scripts/evaluate.py --paper 02          # evaluate one paper by number
  python scripts/evaluate.py --all               # evaluate all training papers in memory/
  python scripts/evaluate.py --paper shani_feinstein_2022   # evaluate by project name
"""
import argparse
import json
import os
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from dotenv import load_dotenv
load_dotenv()

from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

PAPERS_DIR = Path(os.getenv("PAPERS_DIR", "papers"))
MEMORY     = Path("memory")
OUT_DIR    = Path("results/evaluations")
OUT_DIR.mkdir(parents=True, exist_ok=True)

SEV_COLOR = {3: RGBColor(0xC0,0x00,0x00), 2: RGBColor(0xCC,0x66,0x00), 1: RGBColor(0x20,0x6B,0xC0)}
SEV_LABEL = {3: "MUST FIX", 2: "MODERATE", 1: "MINOR"}
ISSUE_KEYS = {"leading_language","double_barreled","social_desirability_bias",
              "ambiguous_wording","scale_appropriateness"}

client = OpenAI()


# ── Helpers ───────────────────────────────────────────────────────────────────
def set_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)

def add_border(table, color="AAAAAA"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    bords = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color); bords.append(el)
    tblPr.append(bords)

def cell_text(cell, text, bold=False, italic=False, color=None, size=10):
    cell.text = ""
    p   = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    if color: run.font.color.rgb = color

def clean_json(raw: str) -> dict:
    s = raw.strip()
    s = re.sub(r"^```json\s*", "", s); s = re.sub(r"^```\s*", "", s)
    s = re.sub(r"\s*```$", "", s)
    return json.loads(s)

def gpt_evaluate(prompt: str) -> str:
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content":
             "You are an expert in marketing survey methodology and consumer behavior research. "
             "Evaluate surveys rigorously and return structured JSON only."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.0,
    )
    return resp.choices[0].message.content


# ── Paper file lookup ─────────────────────────────────────────────────────────
def find_paper_files(project_name: str):
    """
    Find hypotheses.txt for the given project under papers/training/ or papers/test/.
    Returns (hyp_text, paper_label).
    """
    num = project_name.split("_")[0].zfill(2)
    for subdir in ("training", "test"):
        base = PAPERS_DIR / subdir
        if not base.exists():
            continue
        paper_dir = next((d for d in base.iterdir()
                          if d.is_dir() and d.name.split("_")[0].zfill(2) == num), None)
        if not paper_dir:
            continue
        all_hyp = sorted(paper_dir.rglob("hypotheses.txt"))
        if not all_hyp:
            continue
        best_hyp = all_hyp[0]
        for hyp in all_hyp:
            study = hyp.parent.name.lower().replace(" ", "_")
            if project_name.endswith(study):
                best_hyp = hyp
                break
        try:
            hyp_text = best_hyp.read_text(encoding="utf-8", errors="ignore").strip()
        except Exception:
            hyp_text = ""
        study_dir = best_hyp.parent
        paper_label = (f"{paper_dir.name} / {study_dir.name}"
                       if study_dir != paper_dir else paper_dir.name)
        return hyp_text, paper_label
    return "", project_name


# ── Critique normalization ────────────────────────────────────────────────────
def normalize_critique(crit: dict):
    q_crits = crit.get("question_critiques", [])
    issue_map = {}
    for entry in q_crits:
        qid = entry.get("question_id","?")
        if "issues" in entry:
            issues = entry["issues"]
        elif any(k in entry for k in ISSUE_KEYS):
            issues = []
            for k in ISSUE_KEYS:
                if k in entry:
                    sev = entry[k]
                    if isinstance(sev, int) and sev > 0:
                        issues.append({"type": k.replace("_"," "), "severity": sev,
                                       "explanation": entry.get("comments","")})
        else:
            issues = []
        issue_map[qid] = [i for i in issues if isinstance(i, dict)]

    s_raw = crit.get("survey_level_issues", [])
    s_crits = []
    for item in s_raw:
        if isinstance(item, dict):
            if "type" in item:
                s_crits.append(item)
            else:
                for k, v in item.items():
                    if isinstance(v, str) and v:
                        s_crits.append({"type": k.replace("_"," "), "severity": 2, "explanation": v})
        elif isinstance(item, str):
            s_crits.append({"type": "issue", "severity": 2, "explanation": item})
    return issue_map, [i for i in s_crits if isinstance(i, dict)]


# ── Per-paper evaluation ──────────────────────────────────────────────────────
def evaluate_paper(proj: str, mem_path: Path):
    print(f"\n{'─'*60}")
    print(f"  {proj}")

    data  = json.loads(mem_path.read_text())
    v1    = data.get("survey_v1", {})
    v1_qs = v1.get("revised_survey", v1).get("questions", []) if v1 else []
    crit  = data.get("critique", {}) if isinstance(data.get("critique"), dict) else {}
    issue_map, s_crits = normalize_critique(crit)
    all_q_issues  = [i for iss in issue_map.values() for i in iss]
    q_crits_raw   = crit.get("question_critiques", [])

    hyp_text, paper_label = find_paper_files(proj)

    if not v1_qs:
        print("  No V1 questions — skipping.")
        return None

    ai_q_texts = [q.get("question_text","") for q in v1_qs]
    ai_q_ids   = [q.get("question_id","") for q in v1_qs]
    ai_types   = {}
    for q in v1_qs:
        t = str(q.get("input_type","unknown"))
        ai_types[t] = ai_types.get(t,0) + 1

    # ── Metric 1 ──────────────────────────────────────────────────────────────
    print("    [1] Counting questions...")
    metric1 = {
        "ai_question_count":      len(v1_qs),
        "human_question_count":   None,
        "ai_type_distribution":   ai_types,
        "human_type_distribution": None,
        "ground_truth_available": False,
        "ground_truth_source":    None,
    }

    # ── Metric 2: Construct coverage ──────────────────────────────────────────
    print("    [2] Construct coverage...")
    m2_prompt = (
        "You are evaluating an AI-generated survey for a marketing research experiment.\n\n"
        "Research hypotheses:\n" + hyp_text + "\n\n"
        "AI-generated survey questions:\n" + json.dumps(ai_q_texts, indent=2) + "\n\n"
        "Based on the hypotheses, identify 4-6 key constructs that should be measured. "
        "For each construct, score the AI survey's coverage 0-3 (0=missing, 1=partial, 2=adequate, 3=excellent). "
        "Also score these structural elements (0-3): manipulation_checks, attention_checks, demographics.\n\n"
        'Return ONLY JSON: {"key_constructs": [{"name": "...", "ai_score": 0, "human_score": 0, "notes": "..."}], '
        '"structural": {"manipulation_checks": {"ai_score": 0, "human_score": 0, "notes": ""}, '
        '"attention_checks": {"ai_score": 0, "human_score": 0, "notes": ""}, '
        '"demographics": {"ai_score": 0, "human_score": 0, "notes": ""}}, '
        '"overall_coverage_ai": 0, "overall_coverage_human": 0, '
        '"missing_in_ai": [], "extra_in_ai": [], "notes": ""}'
    )
    try:
        metric2 = clean_json(gpt_evaluate(m2_prompt))
    except Exception as e:
        print(f"    Metric 2 parse error: {e}")
        metric2 = {"error": str(e)}

    # ── Metric 3: Wording quality ──────────────────────────────────────────────
    print("    [3] Wording quality...")
    q_list_for_m3 = [{"id": qid, "text": qt} for qid, qt in zip(ai_q_ids, ai_q_texts)]
    m3_prompt = (
        "Score each AI-generated survey question on wording quality (1-5):\n"
        "5=Clear, neutral, standard marketing research phrasing\n"
        "4=Minor issues but acceptable\n"
        "3=Moderate issues (slightly leading, ambiguous, or awkward)\n"
        "2=Significant issues (leading language, double-barreled, confusing)\n"
        "1=Unusable - must be rewritten\n\n"
        "Questions:\n" + json.dumps(q_list_for_m3, indent=2)
        + '\n\nReturn ONLY JSON: {"question_scores": [{"question_id": "...", "score": 5, "issue": null}], '
        '"mean_score": 0.0, "questions_below_3": []}'
    )
    try:
        metric3 = clean_json(gpt_evaluate(m3_prompt))
    except Exception as e:
        print(f"    Metric 3 parse error: {e}")
        metric3 = {"mean_score": "N/A", "question_scores": [], "questions_below_3": []}

    # ── Metric 4: Scale consistency ────────────────────────────────────────────
    print("    [4] Scale consistency...")
    scale_qs = [q for q in v1_qs if str(q.get("input_type","")).lower() in ("scale","likert")]
    scale_ranges = []
    for q in scale_qs:
        cfg = q.get("input_config",{})
        mn  = cfg.get("scale_min", cfg.get("min_value",1))
        mx  = cfg.get("scale_max", cfg.get("scale_points", cfg.get("max_value",7)))
        scale_ranges.append(f"{q.get('question_id','?')}: {mn}–{mx}")
    ends = [r.split("–")[-1] for r in scale_ranges if "–" in r]
    consistent = len(set(ends)) <= 1
    metric4 = {
        "ai_scale_questions":   len(scale_qs),
        "ai_scale_ranges":      scale_ranges,
        "consistent_within_ai": consistent,
    }

    # ── Metric 5: Critic issues ────────────────────────────────────────────────
    print("    [5] Critic issues...")
    sev3_q = [i for i in all_q_issues if i.get("severity")==3]
    sev2_q = [i for i in all_q_issues if i.get("severity")==2]
    sev3_s = [i for i in s_crits if i.get("severity")==3]
    metric5 = {
        "total_question_issues":   len(all_q_issues),
        "severity_3_question":     len(sev3_q),
        "severity_2_question":     len(sev2_q),
        "survey_level_issues":     len(s_crits),
        "survey_level_severity_3": len(sev3_s),
        "issue_rate_per_question": round(len(all_q_issues)/len(v1_qs),2) if v1_qs else 0,
        "severity_3_details":      sev3_q,
        "survey_level_details":    s_crits,
    }

    report = {
        "project":       proj,
        "paper_label":   paper_label,
        "metric1_question_count_and_types": metric1,
        "metric2_construct_coverage":       metric2,
        "metric3_wording_quality":          metric3,
        "metric4_scale_consistency":        metric4,
        "metric5_critic_issues":            metric5,
    }

    (OUT_DIR / f"{proj}_evaluation.json").write_text(json.dumps(report, indent=2))
    write_docx(report, proj, v1_qs, q_crits_raw, issue_map, hyp_text)
    print(f"    Done. Reports: {OUT_DIR / proj}_evaluation.{{json,docx}}")
    return report


# ── Word doc writer ───────────────────────────────────────────────────────────
def write_docx(report, proj, v1_qs, q_crits_raw, issue_map, hyp_text):
    m1 = report["metric1_question_count_and_types"]
    m2 = report["metric2_construct_coverage"]
    m3 = report["metric3_wording_quality"]
    m4 = report["metric4_scale_consistency"]
    m5 = report["metric5_critic_issues"]

    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1.0)

    h = doc.add_heading(f"Evaluation Report: {proj}", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1F,0x38,0x64)
    doc.add_paragraph(f"Paper: {report['paper_label']}").runs[0].italic = True
    doc.add_paragraph()

    doc.add_heading("Research Hypotheses", level=2)
    if hyp_text:
        for line in hyp_text.splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())
    else:
        doc.add_paragraph("Not found.").runs[0].italic = True
    doc.add_paragraph()

    # Metric 1
    doc.add_heading("Metric 1: Question Count & Type Distribution", level=2)
    rows = [("Total questions", str(m1["ai_question_count"]))]
    for t in sorted(m1["ai_type_distribution"].keys()):
        rows.append((t, str(m1["ai_type_distribution"].get(t,0))))
    tbl = doc.add_table(rows=len(rows)+1, cols=2)
    tbl.style = "Table Grid"; add_border(tbl)
    cell_text(tbl.rows[0].cells[0], "Question type", bold=True)
    cell_text(tbl.rows[0].cells[1], "AI count", bold=True)
    for i, (k, v) in enumerate(rows):
        cell_text(tbl.rows[i+1].cells[0], k)
        cell_text(tbl.rows[i+1].cells[1], v)
    doc.add_paragraph()

    # Metric 2
    doc.add_heading("Metric 2: Construct Coverage", level=2)
    cov_ai = m2.get("overall_coverage_ai","N/A")
    p = doc.add_paragraph()
    p.add_run(f"Overall AI coverage: ").bold = True
    p.add_run(f"{cov_ai}%")
    constructs = m2.get("key_constructs", [])
    structural  = m2.get("structural", {})
    all_rows = constructs + [
        {"name": f"[{k.replace('_',' ').title()}]",
         "ai_score": v.get("ai_score","?"), "human_score": v.get("human_score","?"),
         "notes": v.get("notes","")}
        for k, v in structural.items() if isinstance(v, dict)
    ]
    if all_rows:
        tbl2 = doc.add_table(rows=len(all_rows)+1, cols=3)
        tbl2.style = "Table Grid"; add_border(tbl2)
        for ci, hd in enumerate(["Construct", "AI Score (0–3)", "Notes"]):
            set_bg(tbl2.rows[0].cells[ci], "D9E1F2")
            cell_text(tbl2.rows[0].cells[ci], hd, bold=True)
        for i, row in enumerate(all_rows):
            tbl2.rows[i+1].cells[0].text = str(row.get("name",""))
            tbl2.rows[i+1].cells[1].text = str(row.get("ai_score","?"))
            tbl2.rows[i+1].cells[2].text = str(row.get("notes",""))
    missing = m2.get("missing_in_ai", [])
    if missing:
        p = doc.add_paragraph()
        p.add_run("Missing constructs in AI: ").bold = True
        p.add_run(", ".join(missing))
    doc.add_paragraph()

    # Metric 3
    doc.add_heading("Metric 3: Wording Quality (1–5 per question)", level=2)
    p = doc.add_paragraph()
    p.add_run(f"Mean score: ").bold = True
    p.add_run(f"{m3.get('mean_score','N/A')} / 5")
    q_scores = m3.get("question_scores", [])
    if q_scores:
        tbl3 = doc.add_table(rows=len(q_scores)+1, cols=3)
        tbl3.style = "Table Grid"; add_border(tbl3)
        for ci, hd in enumerate(["Question ID", "Score", "Issue"]):
            set_bg(tbl3.rows[0].cells[ci], "D9E1F2")
            cell_text(tbl3.rows[0].cells[ci], hd, bold=True)
        for i, qs in enumerate(q_scores):
            score = qs.get("score","?")
            color = None
            if isinstance(score, (int, float)):
                if score < 3:   color = RGBColor(0xC0,0x00,0x00)
                elif score < 4: color = RGBColor(0xCC,0x66,0x00)
            tbl3.rows[i+1].cells[0].text = str(qs.get("question_id","?"))
            tbl3.rows[i+1].cells[1].text = str(score)
            tbl3.rows[i+1].cells[2].text = str(qs.get("issue") or "—")
            if color:
                for ci in range(3):
                    for run in tbl3.rows[i+1].cells[ci].paragraphs[0].runs:
                        run.font.color.rgb = color
    below3 = m3.get("questions_below_3", [])
    if below3:
        labels = [str(x.get("question_id") or x.get("id") or x) if isinstance(x, dict) else str(x) for x in below3]
        p = doc.add_paragraph()
        p.add_run("Questions needing revision: ").bold = True
        p.add_run(", ".join(labels)).font.color.rgb = RGBColor(0xC0,0x00,0x00)
    doc.add_paragraph()

    # Metric 4
    doc.add_heading("Metric 4: Scale Consistency", level=2)
    rows4 = [
        ("Scale questions (AI)", str(m4["ai_scale_questions"])),
        ("Scale ranges (AI)", ", ".join(m4["ai_scale_ranges"]) or "none"),
        ("Consistent within AI survey", "Yes" if m4["consistent_within_ai"] else "No"),
    ]
    tbl4 = doc.add_table(rows=len(rows4), cols=2)
    tbl4.style = "Table Grid"; add_border(tbl4)
    for i, (k, v) in enumerate(rows4):
        set_bg(tbl4.rows[i].cells[0], "EBF3FB")
        cell_text(tbl4.rows[i].cells[0], k, bold=True)
        cell_text(tbl4.rows[i].cells[1], v)
    doc.add_paragraph()

    # Metric 5
    doc.add_heading("Metric 5: Critic Agent Issues", level=2)
    rows5 = [
        ("Question-level issues (total)",     str(m5["total_question_issues"])),
        ("Severity 3 — MUST FIX",             str(m5["severity_3_question"])),
        ("Severity 2 — MODERATE",             str(m5["severity_2_question"])),
        ("Survey-level issues (total)",        str(m5["survey_level_issues"])),
        ("Survey-level severity 3",            str(m5["survey_level_severity_3"])),
        ("Issue rate per question",            str(m5["issue_rate_per_question"])),
    ]
    tbl5 = doc.add_table(rows=len(rows5)+1, cols=2)
    tbl5.style = "Table Grid"; add_border(tbl5)
    set_bg(tbl5.rows[0].cells[0], "D9E1F2"); set_bg(tbl5.rows[0].cells[1], "D9E1F2")
    cell_text(tbl5.rows[0].cells[0], "Metric", bold=True)
    cell_text(tbl5.rows[0].cells[1], "Count",  bold=True)
    for i, (k, v) in enumerate(rows5):
        cell_text(tbl5.rows[i+1].cells[0], k)
        cell_text(tbl5.rows[i+1].cells[1], v)
    doc.add_paragraph()

    doc.save(OUT_DIR / f"{proj}_evaluation.docx")


# ── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Evaluate AI survey(s) against ground truth")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--paper", type=str,
                       help="Paper number (e.g. 02) or project name (e.g. shani_feinstein_2022)")
    group.add_argument("--all", action="store_true",
                       help="Evaluate all papers found in memory/ (excluding shani_feinstein_2022 and hotel_booking_study)")
    args = parser.parse_args()

    skip = {"hotel_booking_study", "shani_feinstein_2022"}

    if args.all:
        files = sorted(MEMORY.glob("*.json"))
        latest = {}
        for f in files:
            proj = "_".join(f.stem.split("_")[:-2])
            if proj not in skip and not any(s in proj for s in skip):
                if proj not in latest or f.stat().st_mtime > latest[proj].stat().st_mtime:
                    latest[proj] = f
        print(f"Evaluating {len(latest)} papers...")
        results = {}; failed = []
        for proj, path in sorted(latest.items()):
            try:
                r = evaluate_paper(proj, path)
                if r: results[proj] = r
            except Exception as e:
                print(f"  FAILED {proj}: {e}")
                failed.append(proj)
        print(f"\n{'='*60}")
        print(f"Completed: {len(results)}  |  Failed: {len(failed)}")
        if failed: print("Failed:", failed)
        print(f"Reports saved to: {OUT_DIR}/")

    else:
        # Find by number or name
        query = args.paper.strip()
        files = sorted(MEMORY.glob("*.json"))
        latest = {}
        for f in files:
            proj = "_".join(f.stem.split("_")[:-2])
            if proj not in latest or f.stat().st_mtime > latest[proj].stat().st_mtime:
                latest[proj] = f

        match = None
        # Try exact match first
        if query in latest:
            match = (query, latest[query])
        else:
            # Try matching by leading number
            query_num = query.zfill(2)
            for proj, path in latest.items():
                if proj.split("_")[0].zfill(2) == query_num:
                    match = (proj, path)
                    break

        if not match:
            print(f"No memory file found for paper '{query}'.")
            print(f"Available projects: {sorted(latest.keys())}")
            sys.exit(1)

        proj, path = match
        evaluate_paper(proj, path)
