"""
export/before_after.py
Produces before_after_survey.docx — two clean survey printouts:
  Section 1: BEFORE  (v0 — raw converted survey)
  Section 2: AFTER   (final — post-critique revision)
             New questions are shaded green. Critic issues inlined under each question.

Reads from: results/shani_feinstein_2022/ and memory/shani_feinstein_2022_*.json

Usage:
  python export/before_after.py
  python export/before_after.py --project my_project_name
"""
import argparse
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DEFAULT_PROJECT = "shani_feinstein_2022"

SEV_COLOR = {3: RGBColor(0xC0,0x00,0x00), 2: RGBColor(0xCC,0x66,0x00), 1: RGBColor(0x20,0x6B,0xC0)}
SEV_LABEL = {3: "MUST FIX", 2: "MODERATE", 1: "MINOR"}


def set_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_border(table, color="AAAAAA"):
    tbl   = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    bords = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        bords.append(el)
    tblPr.append(bords)


def render_question(doc, q, index, is_new=False, show_issues=False, issue_map=None):
    qid   = q["question_id"]
    qtext = q["question_text"]
    qtype = q["input_type"]
    cfg   = q.get("input_config", {})
    issue_map = issue_map or {}

    p = doc.add_paragraph()
    num = p.add_run(f"Q{index}.  ")
    num.bold = True
    num.font.size = Pt(11)
    if is_new:
        num.font.color.rgb = RGBColor(0x1F, 0x75, 0x23)
    body = p.add_run(qtext)
    body.font.size = Pt(11)
    if is_new:
        body.bold = True

    if is_new:
        tag = doc.add_paragraph()
        run = tag.add_run("  ★  NEW — added by critic-driven revision pass")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x75, 0x23)

    if qtype == "scale":
        labels = cfg.get("scale_labels", [])
        mn     = cfg.get("scale_min", 1)
        if labels:
            n = len(labels)
            tbl = doc.add_table(rows=2, cols=n)
            tbl.style = "Table Grid"
            add_border(tbl, "CCCCCC")
            for j, lbl in enumerate(labels):
                nc = tbl.rows[0].cells[j]
                nc.text = str(mn + j)
                nc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                nc.paragraphs[0].runs[0].font.size = Pt(9)
                nc.paragraphs[0].runs[0].bold = True
                set_bg(nc, "EBF3FB" if not is_new else "EAF4EA")
                lc = tbl.rows[1].cells[j]
                lc.text = lbl
                lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                lc.paragraphs[0].runs[0].font.size = Pt(8)
            doc.add_paragraph()
    elif qtype == "multiple_choice":
        for opt in cfg.get("options", []):
            op = doc.add_paragraph(style="List Bullet")
            op.add_run(opt).font.size = Pt(10)
    elif qtype in ("open_text", "text_input"):
        op = doc.add_paragraph()
        r  = op.add_run("[ Open response field ]")
        r.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        r.font.size = Pt(10)
        ml = cfg.get("max_length")
        if ml:
            op.add_run(f"  (max {ml} characters)").font.size = Pt(9)
    elif qtype == "matrix":
        stmts  = cfg.get("statements", [])
        labels = cfg.get("scale_labels", [])
        if stmts and labels:
            ncols = len(labels) + 1
            tbl   = doc.add_table(rows=len(stmts)+1, cols=ncols)
            tbl.style = "Table Grid"
            add_border(tbl, "CCCCCC")
            set_bg(tbl.rows[0].cells[0], "F2F2F2")
            for j, lbl in enumerate(labels):
                c = tbl.rows[0].cells[j+1]
                c.text = lbl
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                c.paragraphs[0].runs[0].font.size = Pt(8)
                c.paragraphs[0].runs[0].bold = True
                set_bg(c, "EBF3FB" if not is_new else "EAF4EA")
            for r_i, stmt in enumerate(stmts):
                tbl.rows[r_i+1].cells[0].text = stmt
                tbl.rows[r_i+1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                for j in range(len(labels)):
                    c = tbl.rows[r_i+1].cells[j+1]
                    c.text = "○"
                    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

    if show_issues:
        issues = issue_map.get(qid, [])
        if issues:
            for iss in issues:
                sev   = iss.get("severity", 0)
                itype = iss.get("type", "?")
                expl  = iss.get("explanation", "")
                color = SEV_COLOR.get(sev, RGBColor(0,0,0))
                label = SEV_LABEL.get(sev, f"SEV {sev}")
                ip    = doc.add_paragraph(style="List Bullet")
                ip.paragraph_format.left_indent = Inches(0.3)
                tag   = ip.add_run(f"[CRITIC — {label}]  ")
                tag.bold = True
                tag.font.color.rgb = color
                tag.font.size = Pt(9)
                typ   = ip.add_run(f"{itype}:  ")
                typ.bold = True
                typ.font.size = Pt(9)
                ip.add_run(expl).font.size = Pt(9)

    doc.add_paragraph()


def build_doc(project_name: str):
    dir_path  = Path("results") / project_name
    mem_files = sorted(Path("memory").glob(f"{project_name}_*.json"))
    if not mem_files:
        raise FileNotFoundError(f"No memory file found for project '{project_name}'")

    mem         = json.loads(mem_files[-1].read_text())
    v0_qs       = mem["survey_v0"].get("revised_survey", mem["survey_v0"]).get("questions", [])
    final_data  = json.loads((dir_path / "ai_survey.json").read_text())
    final_qs    = final_data.get("revised_survey", final_data).get("questions", [])
    critique    = json.loads((dir_path / "critique.json").read_text())

    v0_ids = {q["question_id"] for q in v0_qs}

    issue_map = {}
    for entry in critique.get("question_critiques", []):
        qid    = entry.get("question_id", "?")
        issues = entry.get("issues", [])
        if not issues and "severity" in entry:
            issues = [entry]
        issue_map[qid] = issues
    survey_issues = critique.get("survey_level_issues", [])

    doc = Document()
    for section in doc.sections:
        section.left_margin  = Inches(1.0)
        section.right_margin = Inches(1.0)

    # SECTION 1 — BEFORE
    h1 = doc.add_heading("BEFORE  —  Original Converted Survey", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    meta = doc.add_paragraph()
    meta.add_run("State: ").bold = True
    meta.add_run("Raw output from the Convert agent — no enhancements or critic pass applied yet.")
    meta.add_run(f"\nQuestion count: ").bold = True
    meta.add_run(str(len(v0_qs)))
    doc.add_paragraph()
    for idx, q in enumerate(v0_qs, 1):
        render_question(doc, q, idx, is_new=False, show_issues=False, issue_map=issue_map)
        doc.add_paragraph("─" * 72).runs[0].font.color.rgb = RGBColor(0xDD,0xDD,0xDD)
        doc.add_paragraph()

    # SECTION 2 — AFTER
    doc.add_page_break()
    h2 = doc.add_heading("AFTER  —  Final Survey (Post-Critique Revision)", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x1F, 0x75, 0x23)
    meta2 = doc.add_paragraph()
    meta2.add_run("State: ").bold = True
    meta2.add_run(
        "Output after parallel enhancement + critic review + Step 4 revision pass. "
        "Questions marked ★ NEW were added by the agent to address critic issues."
    )
    meta2.add_run(f"\nQuestion count: ").bold = True
    meta2.add_run(f"{len(final_qs)}  ({len(final_qs)-len(v0_qs)} added)")
    doc.add_paragraph()

    leg = doc.add_paragraph()
    leg.add_run("Legend:  ").bold = True
    for sev, label in SEV_LABEL.items():
        r = leg.add_run(f"■ {label}  ")
        r.font.color.rgb = SEV_COLOR[sev]
        r.bold = True
    doc.add_paragraph()

    for idx, q in enumerate(final_qs, 1):
        is_new = q["question_id"] not in v0_ids
        render_question(doc, q, idx, is_new=is_new, show_issues=True, issue_map=issue_map)
        doc.add_paragraph("─" * 72).runs[0].font.color.rgb = RGBColor(0xDD,0xDD,0xDD)
        doc.add_paragraph()

    # Survey-level issues appendix
    doc.add_page_break()
    doc.add_heading("Appendix — Survey-Level Critic Issues", level=1)
    doc.add_paragraph(
        "These issues apply to the survey as a whole (not tied to a single question). "
        "The Step 4 revision pass addressed severity-3 items by adding new questions."
    )
    doc.add_paragraph()
    for iss in survey_issues:
        sev   = iss.get("severity", 0)
        itype = iss.get("type", "?").replace("_", " ").title()
        expl  = iss.get("explanation", "")
        color = SEV_COLOR.get(sev, RGBColor(0,0,0))
        label = SEV_LABEL.get(sev, f"SEV {sev}")
        p   = doc.add_paragraph()
        tag = p.add_run(f"[{label}]  {itype}\n")
        tag.bold = True
        tag.font.color.rgb = color
        p.add_run(expl).font.size = Pt(10)
        doc.add_paragraph()

    doc.add_paragraph("Generated by Survey Designer Agent").runs[0].italic = True

    out = dir_path / "before_after_survey.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Export before/after survey comparison")
    parser.add_argument("--project", default=DEFAULT_PROJECT,
                        help=f"Project name (default: {DEFAULT_PROJECT})")
    args = parser.parse_args()
    build_doc(args.project)
